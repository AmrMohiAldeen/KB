from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Table_Architecture_and_Behavior_Reference.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "172B4D"
MUTED = "5B677A"
LIGHT_BLUE = "E8EEF5"
LIGHTER_BLUE = "F4F7FA"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "D7DEE8"
WHITE = "FFFFFF"
GREEN = "EAF5EE"
GREEN_TEXT = "276749"
GOLD = "FFF4D6"
GOLD_TEXT = "7A5A00"
RED = "FBECEC"
RED_TEXT = "9B1C1C"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_TOP_BOTTOM = 80
CELL_MARGIN_START_END = 120


def set_run_font(
    run,
    *,
    name: str = "Calibri",
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    return run


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (
        ("top", CELL_MARGIN_TOP_BOTTOM),
        ("bottom", CELL_MARGIN_TOP_BOTTOM),
        ("start", CELL_MARGIN_START_END),
        ("end", CELL_MARGIN_START_END),
    ):
        child = tc_mar.find(qn(f"w:{tag}"))
        if child is None:
            child = OxmlElement(f"w:{tag}")
            tc_mar.append(child)
        child.set(qn("w:w"), str(value))
        child.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]):
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError(f"Table widths must sum to {CONTENT_WIDTH_DXA}: {widths_dxa}")

    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_keep_with_next(paragraph, value=True):
    paragraph.paragraph_format.keep_with_next = value


def set_paragraph_box(paragraph, fill: str, border_color: str):
    p_pr = paragraph._p.get_or_add_pPr()
    shading = p_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        p_pr.append(shading)
    shading.set(qn("w:fill"), fill)

    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        border = borders.find(qn(f"w:{edge}"))
        if border is None:
            border = OxmlElement(f"w:{edge}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "8")
        border.set(qn("w:space"), "6")
        border.set(qn("w:color"), border_color)


def style_paragraph(
    paragraph,
    *,
    before=0,
    after=6,
    line=1.25,
    keep_next=False,
    keep_together=False,
):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.keep_with_next = keep_next
    fmt.keep_together = keep_together


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr)
    run._r.append(fld_char_2)
    set_run_font(run, size=9, color=MUTED)


def add_header_footer(document: Document):
    section = document.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style_paragraph(p, after=0, line=1.0)
    set_run_font(
        p.add_run("KB Frontend Table Architecture"),
        size=9,
        color=MUTED,
        bold=True,
    )
    set_run_font(
        p.add_run("  |  Engineering Reference"),
        size=9,
        color=MUTED,
    )

    footer = section.footer
    p = footer.paragraphs[0]
    add_page_number(p)


def configure_document(document: Document):
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.line_spacing = 1.1

    code = styles.add_style("Code Block", 1)
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    code._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    code.font.size = Pt(9)
    code.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    code.paragraph_format.space_before = Pt(3)
    code.paragraph_format.space_after = Pt(7)
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.right_indent = Inches(0.18)
    code.paragraph_format.line_spacing = 1.05

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.font.italic = True
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(4)

    add_header_footer(document)


def add_title_page(document: Document):
    p = document.add_paragraph()
    style_paragraph(p, before=68, after=8, line=1.0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(
        p.add_run("KB FRONTEND"),
        size=11,
        color=BLUE,
        bold=True,
    )

    p = document.add_paragraph()
    style_paragraph(p, after=9, line=1.0, keep_together=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(
        p.add_run("Table Architecture and Behavior Reference"),
        size=28,
        color=INK,
        bold=True,
    )

    p = document.add_paragraph()
    style_paragraph(p, after=26, line=1.15)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(
        p.add_run(
            "A detailed guide to table extensions, commands, toolbar state, "
            "headers, borders, resizing, dragging, selection, and production hardening"
        ),
        size=13,
        color=MUTED,
        italic=True,
    )

    table = document.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2700, 6660])
    metadata = [
        ("Repository", r"C:\gamalearn\kb-frontend"),
        ("Primary stack", "Next.js 16.2.7, React 19.2.4, Tiptap 3.26.0, ProseMirror tables"),
        ("Document date", "June 10, 2026"),
        ("Scope", "All table-related extensions, commands, plugins, toolbar UI, utilities, styles, and tests"),
        ("Verification", "40 Vitest tests, 10 Playwright tests, lint, TypeScript, browser inspection, production build"),
    ]
    for row, (label, value) in zip(table.rows, metadata):
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        set_run_font(row.cells[0].paragraphs[0].add_run(label), bold=True, color=DARK_BLUE)
        set_run_font(row.cells[1].paragraphs[0].add_run(value), color=INK)

    p = document.add_paragraph()
    style_paragraph(p, before=22, after=6, line=1.15)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(
        p.add_run("Design intent"),
        size=10,
        color=BLUE,
        bold=True,
    )
    p = document.add_paragraph()
    style_paragraph(p, after=0, line=1.25)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(
        p.add_run(
            "The implementation treats table state as editor data, DOM geometry as a temporary "
            "interaction surface, and every command boundary as potentially invalid."
        ),
        size=10.5,
        color=MUTED,
    )
    document.add_page_break()


def add_heading(document: Document, text: str, level: int = 1):
    return document.add_heading(text, level=level)


def add_para(
    document: Document,
    text: str = "",
    *,
    bold_label: str | None = None,
    style: str | None = None,
    color: str = INK,
    after: int = 6,
    line: float = 1.25,
    keep_together: bool = False,
):
    p = document.add_paragraph(style=style)
    style_paragraph(p, after=after, line=line, keep_together=keep_together)
    if bold_label and text.startswith(bold_label):
        set_run_font(p.add_run(bold_label), bold=True, color=DARK_BLUE)
        set_run_font(p.add_run(text[len(bold_label) :]), color=color)
    else:
        set_run_font(p.add_run(text), color=color)
    return p


def add_bullets(document: Document, items: list[str]):
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        style_paragraph(p, after=4, line=1.25)
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        set_run_font(p.add_run(item), color=INK)


def add_numbered(document: Document, items: list[str]):
    for item in items:
        p = document.add_paragraph(style="List Number")
        style_paragraph(p, after=4, line=1.25)
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        set_run_font(p.add_run(item), color=INK)


def add_callout(document: Document, title: str, body: str, kind="info"):
    palette = {
        "info": (LIGHTER_BLUE, DARK_BLUE),
        "success": (GREEN, GREEN_TEXT),
        "warning": (GOLD, GOLD_TEXT),
        "risk": (RED, RED_TEXT),
    }
    fill, title_color = palette[kind]
    p = document.add_paragraph()
    set_paragraph_box(p, fill, title_color)
    style_paragraph(p, before=2, after=8, line=1.15, keep_together=True)
    set_run_font(p.add_run(f"{title}: "), bold=True, color=title_color)
    set_run_font(p.add_run(body), color=INK)


def add_code(document: Document, code: str):
    p = document.add_paragraph()
    p.style = document.styles["Code Block"]
    set_paragraph_box(p, LIGHTER_BLUE, BLUE)
    style_paragraph(p, before=2, after=8, line=1.1, keep_together=True)
    set_run_font(p.add_run(code), name="Consolas", size=9, color=DARK_BLUE)


def add_table(
    document: Document,
    headers: list[str],
    rows: list[list[str]],
    widths: list[int],
    *,
    header_fill=LIGHT_BLUE,
    first_col_fill: str | None = None,
    font_size=9.5,
):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        style_paragraph(p, after=0, line=1.1)
        set_run_font(p.add_run(header), size=font_size, color=DARK_BLUE, bold=True)
    repeat_table_header(table.rows[0])

    for row_values in rows:
        row = table.add_row()
        for i, value in enumerate(row_values):
            cell = row.cells[i]
            if i == 0 and first_col_fill:
                set_cell_shading(cell, first_col_fill)
            p = cell.paragraphs[0]
            style_paragraph(p, after=0, line=1.12)
            set_run_font(
                p.add_run(value),
                size=font_size,
                color=DARK_BLUE if i == 0 else INK,
                bold=(i == 0 and first_col_fill is not None),
            )
    set_table_geometry(table, widths)
    add_para(document, "", after=3)
    return table


def add_file_reference(
    document: Document,
    path: str,
    responsibility: str,
    behavior: str,
    dependencies: str,
    notes: str,
):
    add_heading(document, path, level=3)
    add_table(
        document,
        ["Aspect", "Details"],
        [
            ["Responsibility", responsibility],
            ["Key behavior", behavior],
            ["Important dependencies", dependencies],
            ["Maintenance notes", notes],
        ],
        [1900, 7460],
        first_col_fill=LIGHTER_BLUE,
        font_size=9.3,
    )


def build_document():
    document = Document()
    configure_document(document)
    add_title_page(document)

    add_heading(document, "Contents", level=1)
    contents = [
        "1. Executive architecture overview",
        "2. Runtime architecture and data flow",
        "3. Extension configuration and persisted schema",
        "4. Command architecture and safe failure model",
        "5. Header row and header column behavior",
        "6. Selection, deletion, merge, and split behavior",
        "7. Toolbar architecture and state synchronization",
        "8. Resizing and dimensions",
        "9. Borders and rendered styling",
        "10. Table dragging, positioning, and reordering",
        "11. DOM and interaction utilities",
        "12. File-by-file responsibility reference",
        "13. Edge-case behavior matrix",
        "14. Hardening improvements",
        "15. Test and verification strategy",
        "16. Remaining risks and production concerns",
        "17. Safe extension and maintenance guide",
    ]
    add_numbered(document, [item.split(". ", 1)[1] for item in contents])
    add_callout(
        document,
        "Reading guide",
        "Sections 1-11 explain the system by behavior. Section 12 is the detailed file map. "
        "Sections 13-17 focus on edge cases, hardening, verification, and future maintenance.",
        "info",
    )
    document.add_page_break()

    add_heading(document, "1. Executive Architecture Overview", level=1)
    add_para(
        document,
        "The table implementation is a layered Tiptap/ProseMirror subsystem. The editor schema "
        "stores durable table data; commands make validated document mutations; custom plugins "
        "manage geometry-heavy mouse interactions; React toolbar components expose the supported "
        "operations; and CSS renders persistent attributes and temporary interaction handles.",
    )
    add_callout(
        document,
        "Core invariant",
        "Persistent state belongs on ProseMirror nodes. Temporary preview state belongs in the DOM "
        "or plugin-local session state. A preview is restored before the final node mutation is committed.",
        "success",
    )
    add_heading(document, "Architecture at a glance", level=2)
    add_code(
        document,
        "EditorWorkspace / KnowledgeBaseEditor\n"
        "        |\n"
        "        v\n"
        "EditorToolbar + TableControls -----> safe command boundary\n"
        "        |                                  |\n"
        "        v                                  v\n"
        "Tiptap table extensions ----------> ProseMirror transactions\n"
        "        |                                  |\n"
        "        v                                  v\n"
        "Custom table node attributes       document history / undo-redo\n"
        "        |\n"
        "        +--> TableView applies width, offset, and borders to DOM\n"
        "        +--> Plugins preview and commit row resize, outer resize, and drag\n"
        "        +--> CSS renders cells, borders, selections, and handles"
    )
    add_heading(document, "Design goals", level=2)
    add_bullets(
        document,
        [
            "Keep table behavior modular and colocated under src/features/editor/table.",
            "Preserve semantic header intent after structural changes, including first-row and first-column mutations.",
            "Return false instead of throwing when editor, selection, table map, DOM, or plugin state is invalid.",
            "Make geometry changes undoable as single history actions.",
            "Keep toolbar enabled/disabled state consistent with the action that will actually run.",
            "Persist border, width, offset, and row-height choices through JSON and HTML serialization.",
        ],
    )

    add_heading(document, "2. Runtime Architecture and Data Flow", level=1)
    add_heading(document, "2.1 Editor startup", level=2)
    add_numbered(
        document,
        [
            "KnowledgeBaseEditor creates a client-side Tiptap Editor with getEditorExtensions().",
            "getEditorExtensions() registers the standard editor extensions, TableKit with selected defaults disabled, and the custom tableExtensions array.",
            "The custom KnowledgeBaseTable replaces the default table node, while custom cell and header extensions add persisted rowHeight attributes.",
            "RowResizePlugin, TableOuterResizePlugin, and TableDragHandlePlugin are wrapped as Tiptap extensions and installed as ProseMirror plugins.",
            "EditorToolbar renders the global table creation picker and the context-sensitive TableControls component.",
        ],
    )
    add_heading(document, "2.2 Mutation lifecycle", level=2)
    add_code(
        document,
        "User action\n"
        "  -> toolbar / keyboard / plugin event\n"
        "  -> validate editor + active table + selection\n"
        "  -> create or run ProseMirror transaction\n"
        "  -> normalize headers or attributes when required\n"
        "  -> dispatch transaction\n"
        "  -> history captures one undoable action\n"
        "  -> TableView and React toolbar re-read the new editor state"
    )
    add_heading(document, "2.3 Persistent versus temporary state", level=2)
    add_table(
        document,
        ["State category", "Stored in", "Examples", "Why"],
        [
            [
                "Persistent document state",
                "ProseMirror node attributes",
                "tableWidthPct, tableOffsetPct, border flags, rowHeight",
                "Serializes to JSON/HTML and participates in undo/redo.",
            ],
            [
                "Plugin state",
                "PluginKey state",
                "Active row resize target, active outer resize table",
                "Tracks interaction targets and maps positions through transactions.",
            ],
            [
                "Drag-session state",
                "Plugin closure / MouseDragSession",
                "Start coordinates, preview values, dominant drag axis",
                "Exists only for the current pointer interaction.",
            ],
            [
                "DOM preview state",
                "Inline styles, data attributes, temporary style element",
                "Live row height, live outer width, live horizontal offset",
                "Gives immediate feedback without producing many history entries.",
            ],
            [
                "React toolbar state",
                "useEditorState selector result",
                "Visibility, header flags, border flags, command capabilities",
                "Keeps controls synchronized with current selection and editor state.",
            ],
        ],
        [1650, 1800, 2650, 3260],
        font_size=8.9,
    )

    add_heading(document, "3. Extension Configuration and Persisted Schema", level=1)
    add_heading(document, "3.1 TableKit replacement strategy", level=2)
    add_para(
        document,
        "The editor still uses Tiptap's TableKit to provide the table row extension and common "
        "table infrastructure, but it explicitly disables the default table, tableCell, and "
        "tableHeader nodes. Those three nodes are replaced by custom extensions so the application "
        "can own attributes, rendering, keyboard behavior, and plugin configuration.",
    )
    add_code(
        document,
        "TableKit.configure({ table: false, tableCell: false, tableHeader: false })\n"
        "...tableExtensions"
    )
    add_heading(document, "3.2 Table extension options", level=2)
    add_table(
        document,
        ["Option", "Configured value", "Effect"],
        [
            ["resizable", "true", "Enables Tiptap's internal column-resizing plugin and TableView."],
            ["cellMinWidth", "40", "Sets the minimum internal column/cell width to 40 pixels."],
            ["lastColumnResizable", "false", "Prevents the final internal column edge from acting as the outer-table resize edge."],
            ["draggable", "true", "Marks the table node as draggable; the custom drag handle controls the interaction."],
            ["View", "KnowledgeBaseTableView", "Applies stored width, offset, and border attributes whenever the table node view is created or updated."],
        ],
        [1900, 1800, 5660],
        font_size=9.2,
    )
    add_heading(document, "3.3 Persisted table attributes", level=2)
    add_table(
        document,
        ["Attribute", "Default", "Rendered form", "Meaning / constraints"],
        [
            ["tableWidthPct", "100", "data-table-width-pct and width/CSS variable", "Outer table width, clamped to 10-100 percent."],
            ["tableOffsetPct", "0", "data-table-offset-pct and margin-left/CSS variable", "Horizontal offset, clamped so width + offset never exceeds 100 percent."],
            ["borderTopEnabled", "true", "data-table-border-top", "Controls the visible top edge."],
            ["borderRightEnabled", "true", "data-table-border-right", "Controls the visible right edge."],
            ["borderBottomEnabled", "true", "data-table-border-bottom", "Controls the visible bottom edge."],
            ["borderLeftEnabled", "true", "data-table-border-left", "Controls the visible left edge."],
            ["borderInnerEnabled", "true", "data-table-border-inner", "Controls internal grid lines while allowing outer edges to remain visible."],
        ],
        [1800, 900, 2750, 3910],
        font_size=8.8,
    )
    add_heading(document, "3.4 Persisted cell attributes", level=2)
    add_para(
        document,
        "Both tableCell and tableHeader nodes receive the same rowHeight attribute. The value is "
        "normalized to a positive integer of at least 20 pixels and rendered as data-row-height "
        "plus an inline height style. Storing the value on every unique cell that participates in "
        "the row allows row height to survive serialization and complex table structures.",
    )
    add_callout(
        document,
        "Compatibility behavior",
        "Tables loaded without the custom border attributes default to all borders enabled. Existing "
        "HTML width and margin-left percentages are accepted as fallback inputs when data attributes are missing.",
        "info",
    )

    add_heading(document, "4. Command Architecture and Safe Failure Model", level=1)
    add_heading(document, "4.1 Why commands are wrapped", level=2)
    add_para(
        document,
        "Direct Tiptap/ProseMirror table commands assume a valid table selection and structurally "
        "valid table map. In a UI, those assumptions can be invalidated by asynchronous React "
        "updates, read-only mode changes, destroyed editors, malformed imported content, stale "
        "positions, or selections that cover an entire table. The command module is the safety "
        "boundary that converts those cases into a false result instead of a runtime exception.",
    )
    add_heading(document, "4.2 Safe command families", level=2)
    add_table(
        document,
        ["Function", "Purpose", "Safety behavior"],
        [
            ["runTableStructureCommand", "Insert/delete rows or columns while preserving header intent.", "Requires usable active table; catches failures; normalizes headers in the same chain."],
            ["canRunTableCommand", "Compute accurate toolbar capability state.", "Returns false for invalid editor/table state and explicitly blocks full-table row/column deletion."],
            ["runTableActionCommand", "Run merge, split, delete-table, or header toggle actions.", "Requires editable active table and catches command failures."],
            ["updateTableBorders", "Persist partial border settings.", "Only updates an active editable table."],
            ["insertTable", "Insert a new table with a header row.", "Validates editor state and positive integer dimensions."],
            ["reorderTableRow", "Move a row while preserving edge header semantics.", "Validates indexes and performs full header normalization after moveTableRow."],
        ],
        [2000, 3350, 4010],
        font_size=8.8,
    )
    add_heading(document, "4.3 Validation sequence", level=2)
    add_numbered(
        document,
        [
            "Reject null or undefined editors.",
            "Reject destroyed or read-only editors.",
            "Resolve the active table from a table NodeSelection or by walking up from the current selection.",
            "Safely construct a TableMap when the operation requires table geometry.",
            "Validate command-specific inputs, such as dimensions, row indexes, or selected rectangle coverage.",
            "Run the underlying command inside a try/catch boundary.",
            "Return false without mutating the document if any prerequisite fails.",
        ],
    )
    add_callout(
        document,
        "Important distinction",
        "canRunTableCommand does more than call editor.can(). ProseMirror can report deletion as "
        "available for a selection covering all rows or all columns even though the actual command "
        "refuses it. The wrapper checks selectedRect explicitly so toolbar state matches runtime behavior.",
        "warning",
    )

    add_heading(document, "5. Header Row and Header Column Behavior", level=1)
    add_heading(document, "5.1 Header intent", level=2)
    add_para(
        document,
        "Header state is treated as an edge-level semantic intent: the first row may be a header row, "
        "the first column may be a header column, or both may be active. The implementation detects "
        "that intent before structural mutation using rowIsHeader and columnIsHeader.",
    )
    add_heading(document, "5.2 Why normalization is necessary", level=2)
    add_para(
        document,
        "Tiptap's structural commands copy neighboring cell types when adding or deleting rows and "
        "columns. That behavior can move or duplicate tableHeader node types in ways that do not "
        "match the application's edge-header model. For example, inserting a row above the first "
        "row may leave the former header row as headers, and deleting the first header column may "
        "fail to promote the new first column.",
    )
    add_heading(document, "5.3 Normalization algorithm", level=2)
    add_numbered(
        document,
        [
            "Snapshot hasHeaderRow and hasHeaderColumn before the structural command runs.",
            "Run the Tiptap add/delete command in a chain.",
            "Resolve the active table from the transaction's updated document.",
            "Use TableMap to visit cells along the first two row/column edges for normal structural changes.",
            "Set each visited unique cell to tableHeader only when it belongs to an intended header edge; otherwise set it to tableCell.",
            "For explicit row reordering, normalize the entire table so a moved former header row is demoted everywhere.",
            "Commit normalization in the same transaction/history action as the structure change.",
        ],
    )
    add_table(
        document,
        ["Operation", "Expected header result"],
        [
            ["Insert row above the header row", "New first row becomes the header row; displaced former header row becomes normal cells except for an active header column."],
            ["Insert row below the header row", "Inserted row is normal except for an active first-column header cell."],
            ["Delete the header row", "The next row is promoted to the header row."],
            ["Insert column before header column", "New first column becomes the header column; displaced former header column is demoted."],
            ["Delete header column", "The new first column is promoted to header cells."],
            ["Reorder first row away from edge", "The new first row is promoted; moved row is demoted except for header-column cells."],
            ["Undo / redo", "Header semantics return with the same structural transaction."],
        ],
        [2800, 6560],
        font_size=9.0,
    )

    add_heading(document, "6. Selection, Deletion, Merge, and Split Behavior", level=1)
    add_heading(document, "6.1 Active table resolution", level=2)
    add_para(
        document,
        "An active table is resolved in two ways. If the selection is a NodeSelection of a table, "
        "the selected table position is used directly. Otherwise, the implementation walks from "
        "selection.$from toward the document root until it finds a node named table. This supports "
        "text selections, cell selections, and whole-table selections.",
    )
    add_heading(document, "6.2 Selected-cell deletion", level=2)
    add_para(
        document,
        "The custom table extension overrides Backspace, Mod-Backspace, Delete, and Mod-Delete. "
        "When a CellSelection is active, deleteCellSelection clears the contents of selected cells "
        "but preserves the table grid. This intentionally differs from Tiptap's default shortcut "
        "that deletes the entire table when every cell is selected.",
    )
    add_callout(
        document,
        "Deletion policy",
        "Keyboard deletion clears cell content. Structural row/column deletion and whole-table "
        "deletion are explicit toolbar actions. This reduces accidental destructive behavior.",
        "success",
    )
    add_heading(document, "6.3 Merge and split", level=2)
    add_para(
        document,
        "Merge and split use Tiptap's table commands through runTableActionCommand. Capability state "
        "is computed with canRunTableCommand, so Merge is enabled only for a valid rectangular "
        "multi-cell CellSelection and Split is enabled only for a merged cell. Integration tests "
        "verify that TableMap remains structurally valid after both actions.",
    )

    add_heading(document, "7. Toolbar Architecture and State Synchronization", level=1)
    add_heading(document, "7.1 Global creation control", level=2)
    add_para(
        document,
        "TableCreationPicker is always available in the main EditorToolbar while the editor is "
        "editable. It offers a 6-by-8 visual grid plus manual dimensions. Manual values are parsed, "
        "clamped to 1-100 rows and 1-20 columns, and passed through the safe insertTable command.",
    )
    add_heading(document, "7.2 Contextual table controls", level=2)
    add_para(
        document,
        "TableControls appears only when an editable editor has an active table. Its state is derived "
        "with useEditorState, which causes React to recompute the selector whenever relevant editor "
        "state changes. This keeps toolbar visibility, active header flags, border checkboxes, and "
        "enabled/disabled command states synchronized with selection changes and document mutations.",
    )
    add_table(
        document,
        ["Toolbar group", "Controls", "Command path"],
        [
            ["Insert", "Row above/below, column before/after", "runTableStructureCommand"],
            ["Cells", "Merge, split", "runTableActionCommand"],
            ["Delete", "Delete row, column, or table", "Structure wrapper for row/column; action wrapper for table"],
            ["Headers", "Toggle header row or column", "runTableActionCommand"],
            ["Borders", "All, outer, inner, and individual edges", "updateTableBorders"],
        ],
        [1600, 3550, 4210],
        font_size=9.0,
    )
    add_heading(document, "7.3 Inactive state", level=2)
    add_para(
        document,
        "The inactive state is a complete typed object with visibility false, disabled capabilities, "
        "no header intent, and default border values. Returning a stable inactive shape makes the "
        "selector predictable and prevents conditional property access from leaking into the component.",
    )

    add_heading(document, "8. Resizing and Dimensions", level=1)
    add_heading(document, "8.1 Internal column resizing", level=2)
    add_para(
        document,
        "Internal column resizing is supplied by Tiptap's resizable table extension. It updates "
        "column widths inside the table but must not change the custom outer table width. Setting "
        "lastColumnResizable to false reserves the table's rightmost boundary for the outer resize plugin.",
    )
    add_heading(document, "8.2 Outer table resizing", level=2)
    add_numbered(
        document,
        [
            "Only an active editable table can expose the outer resize handle.",
            "Hit detection requires the pointer to be within 8 pixels of the right edge and vertically inside the table rectangle.",
            "The drag session records startX, stored width percentage, table position, and container width.",
            "Mouse movement converts horizontal pixels to a percentage delta and applies a live DOM preview.",
            "On mouseup, the preview is restored from stored node attributes, then one setNodeMarkup transaction persists the final width and adjusted offset.",
            "closeHistory makes the completed drag one undoable history action.",
            "Blur, missing DOM, read-only changes, or plugin destruction cancel safely and restore stored state.",
        ],
    )
    add_heading(document, "8.3 Row resizing", level=2)
    add_numbered(
        document,
        [
            "Hit detection activates near the bottom edge of a td or th, or on the explicit row-resize widget.",
            "A temporary style element targets the current table wrapper and row index for a live preview.",
            "The committed height is clamped to at least 20 pixels.",
            "TableMap.positionAt finds every unique cell participating in the row; a visited set prevents duplicate updates for spanning cells.",
            "rowHeight is written to each unique cell in one closeHistory transaction.",
            "Preview style elements and data markers are removed on commit, cancel, blur, read-only transition, or destruction.",
        ],
    )
    add_heading(document, "8.4 Dimension normalization rules", level=2)
    add_table(
        document,
        ["Value", "Rule"],
        [
            ["Outer width", "Finite numeric value, rounded to one decimal place, clamped to 10-100 percent; invalid values become 100."],
            ["Horizontal offset", "Finite numeric value, rounded to one decimal place, clamped to 0 through (100 - width); invalid values become 0."],
            ["Row height", "Finite positive value, rounded to an integer, minimum 20 pixels; invalid values become null or the minimum when clamped."],
            ["Container width", "Must be finite and greater than zero; otherwise plugins use 1 pixel to avoid division by zero."],
        ],
        [2200, 7160],
        font_size=9.2,
    )

    add_heading(document, "9. Borders and Rendered Styling", level=1)
    add_heading(document, "9.1 Border data model", level=2)
    add_para(
        document,
        "Borders are table-level boolean attributes rather than transient wrapper classes. This makes "
        "them durable across JSON serialization, HTML output, editor remounts, undo/redo, and table "
        "node-view updates. Missing values normalize to true for backward compatibility.",
    )
    add_heading(document, "9.2 Rendering strategy", level=2)
    add_para(
        document,
        "KnowledgeBaseTableView applies all border flags as data-table-border-* attributes on the "
        "rendered HTML table. tiptap-table.css uses those attributes to make individual edge colors "
        "transparent, hide all inner borders, and selectively restore enabled outer edges when inner "
        "borders are disabled.",
    )
    add_heading(document, "9.3 Toolbar behavior", level=2)
    add_bullets(
        document,
        [
            "All borders toggles all five flags.",
            "Outer border toggles top, right, bottom, and left together without changing inner borders.",
            "Inner border controls only the internal grid.",
            "Individual edge controls update one attribute.",
            "The Borders toolbar group is active when any border is enabled.",
        ],
    )
    add_callout(
        document,
        "CSS cleanup",
        "Repeated .tiptap and .ProseMirror selectors were consolidated with :is(...). The persistent "
        "data-attribute contract remains unchanged.",
        "info",
    )

    add_heading(document, "10. Table Dragging, Positioning, and Reordering", level=1)
    add_heading(document, "10.1 Drag handle", level=2)
    add_para(
        document,
        "TableDragHandlePlugin creates a button decoration only for the active table. The handle is "
        "hidden and non-draggable in read-only mode. Clicking it creates a table NodeSelection. "
        "Starting a drag serializes the table node to clipboard-compatible HTML and text without "
        "detaching the decoration.",
    )
    add_heading(document, "10.2 Dominant-axis behavior", level=2)
    add_para(
        document,
        "After a small threshold, the drag locks to horizontal or vertical movement. Horizontal "
        "movement adjusts tableOffsetPct. Vertical movement moves the table as a top-level "
        "ProseMirror node. Once chosen, the axis does not change during the drag.",
    )
    add_table(
        document,
        ["Axis", "Preview", "Commit"],
        [
            ["Horizontal", "Applies a live table margin-left/offset preview and repositions the handle.", "Persists tableOffsetPct with setNodeMarkup as one history action."],
            ["Vertical", "Uses the current target block and pointer position relative to its midpoint.", "Deletes and reinserts the entire table node at a valid drop point, then selects it."],
        ],
        [1600, 3900, 3860],
        font_size=9.1,
    )
    add_heading(document, "10.3 Safe vertical placement", level=2)
    add_para(
        document,
        "Vertical placement only moves top-level table nodes. createTableMoveTransaction rejects "
        "invalid positions, nested table positions, drops inside the dragged table, and unsupported "
        "drop points. The drop direction is determined by whether the pointer is above or below the "
        "target block's visual midpoint, with movement direction used only as a fallback.",
    )
    add_heading(document, "10.4 Row reordering API", level=2)
    add_para(
        document,
        "reorderTableRow provides a safe command-level API for moving rows inside a table. It is "
        "currently tested and available to future UI, but no row-reorder toolbar or drag handle is "
        "exposed. The function validates row indexes, uses ProseMirror's moveTableRow, then normalizes "
        "the full table so header intent remains anchored to the first row and first column.",
    )

    add_heading(document, "11. DOM and Interaction Utilities", level=1)
    add_heading(document, "11.1 Table DOM helpers", level=2)
    add_para(
        document,
        "tableDom.ts centralizes all conversions among ProseMirror table positions, node selections, "
        "table wrapper DOM nodes, HTMLTableElement instances, event targets, and editor-relative "
        "overlay positions. Each helper validates inputs and catches view/DOM lookup failures.",
    )
    add_table(
        document,
        ["Helper", "Responsibility"],
        [
            ["getTableNodeAt", "Validate a numeric document position and return only a table node."],
            ["getActiveTable / getActiveTablePos", "Resolve the table around the current selection or table NodeSelection."],
            ["mapTablePos", "Map a stored position through a transaction and reject deleted/non-table results."],
            ["getTableWrapperAtPos / getTableAtPos", "Safely resolve the rendered wrapper and HTML table."],
            ["getOwnerWindow", "Use the editor document's Window for correct instanceof checks and event listeners."],
            ["getClosestHTMLElement", "Safely find event-target ancestors in the correct browser realm."],
            ["requestViewAnimationFrame", "Schedule DOM positioning only while the editor view is alive."],
            ["positionOverlayAtRect", "Convert viewport rectangles to editor-relative overlay geometry."],
        ],
        [2900, 6460],
        font_size=9.0,
    )
    add_heading(document, "11.2 Mouse drag session helper", level=2)
    add_para(
        document,
        "startMouseDragSession owns window-level mousemove, mouseup, and blur listeners. It guarantees "
        "cleanup happens once, distinguishes commit from cancellation, and gives plugins one idempotent "
        "cancel method for read-only transitions and destruction.",
    )
    add_heading(document, "11.3 Row preview helper", level=2)
    add_para(
        document,
        "rowHeightPreview.ts creates a uniquely identified temporary style element and marks the table "
        "wrapper with a matching data attribute. This avoids mutating each cell's inline style during "
        "mousemove and allows the preview to be removed deterministically.",
    )

    document.add_page_break()
    add_heading(document, "12. File-by-File Responsibility Reference", level=1)
    add_para(
        document,
        "This section maps every table-related implementation and test file to its ownership boundary. "
        "Paths are relative to the repository root.",
    )

    add_heading(document, "Implementation Files", level=2)
    file_refs = [
        (
            "src/features/editor/extensions/index.tsx",
            "Assembles the complete Tiptap extension list.",
            "Disables default TableKit table/cell/header nodes and installs the custom tableExtensions.",
            "@tiptap/starter-kit, @tiptap/extension-table, table/extensions",
            "Any new table node replacement must be coordinated here to avoid duplicate extension names.",
        ),
        (
            "src/features/editor/table/extensions/index.ts",
            "Composition root for the table subsystem.",
            "Configures the custom table, row-height cells/headers, and three custom ProseMirror plugins.",
            "KnowledgeBaseTable, TableCellExtensions, RowResizePlugin, TableOuterResizePlugin, TableDragHandlePlugin",
            "Plugin extensions are intentionally small wrappers; keep table-wide configuration here.",
        ),
        (
            "src/features/editor/table/extensions/TableExtension.ts",
            "Owns the custom table node, attributes, node view, and destructive-key behavior.",
            "Persists width/offset/borders; reapplies attributes to DOM; clears selected cell content on Delete/Backspace.",
            "@tiptap/extension-table Table/TableView, tableDimensions, tableBorders",
            "This is the schema contract. Attribute changes require serialization, CSS, and compatibility tests.",
        ),
        (
            "src/features/editor/table/extensions/TableCellExtensions.ts",
            "Adds rowHeight to tableCell and tableHeader.",
            "Normalizes and renders row height consistently for normal and header cells.",
            "@tiptap/extension-table TableCell/TableHeader, tableDimensions",
            "Keep both extensions aligned; rows may contain either node type.",
        ),
        (
            "src/features/editor/table/commands/tableCommands.ts",
            "Central safe command boundary and header normalization layer.",
            "Validates editor/table state, computes capabilities, preserves headers, updates borders, inserts tables, and reorders rows.",
            "Tiptap Editor, ProseMirror TableMap/selectedRect/moveTableRow, tableDom",
            "Toolbar and future UI should call these wrappers rather than raw table commands.",
        ),
        (
            "src/features/editor/table/dom/tableDom.ts",
            "Shared ProseMirror-to-DOM lookup and overlay positioning utilities.",
            "Finds active tables, validates positions, maps positions through transactions, resolves DOM safely.",
            "ProseMirror model/state/view",
            "Use these helpers in plugins to avoid duplicated fragile position walking and cross-window instanceof checks.",
        ),
        (
            "src/features/editor/table/plugins/RowResizePlugin.ts",
            "Interactive row resizing.",
            "Detects row edges, displays a live preview, commits rowHeight to unique row cells, and handles cancellation/read-only state.",
            "TableMap, tableDom, rowHeightPreview, tableDimensions, mouseDragSession",
            "Merged/spanning cells require the visited-cell logic; do not replace it with simple DOM row iteration for persistence.",
        ),
        (
            "src/features/editor/table/plugins/TableOuterResizePlugin.ts",
            "Interactive outer table width resizing.",
            "Detects the right edge, previews percentage width, adjusts offset, and commits one undoable node attribute change.",
            "tableDom, tableDimensions, mouseDragSession, ProseMirror decorations/history",
            "The right-edge vertical-bound check prevents resize activation beside but outside the table.",
        ),
        (
            "src/features/editor/table/plugins/TableDragHandlePlugin.ts",
            "Whole-table dragging and horizontal positioning.",
            "Creates the drag handle, serializes drag payload, locks axis, moves top-level table nodes, or persists horizontal offset.",
            "dropPoint, NodeSelection, tableDom, tableDimensions",
            "Vertical drop placement is based on target midpoint; nested/non-table positions fail safely.",
        ),
        (
            "src/features/editor/table/resizing/tableDimensions.ts",
            "Pure dimension parsing, normalization, clamping, reading, and DOM application.",
            "Enforces width, offset, and row-height invariants and protects against NaN/infinite values.",
            "Browser HTMLTableElement only",
            "Keep geometry math pure and covered by unit tests; plugins should delegate to this module.",
        ),
        (
            "src/features/editor/table/resizing/rowHeightPreview.ts",
            "Temporary live row-height styling.",
            "Creates, updates, and removes unique preview style elements scoped to a table wrapper.",
            "tableDimensions",
            "Preview state must always be restored on cancel, blur, read-only transition, and destruction.",
        ),
        (
            "src/features/editor/table/utils/tableBorders.ts",
            "Border attribute types, defaults, normalization, and DOM application.",
            "Defaults missing borders to enabled and writes data-table-border-* attributes.",
            "HTMLTableElement",
            "This module is the compatibility and type-safety boundary for border flags.",
        ),
        (
            "src/features/editor/table/utils/mouseDragSession.ts",
            "Reusable mouse interaction lifecycle.",
            "Registers movement/commit/cancel listeners and guarantees one-time cleanup.",
            "Window mouse and blur events",
            "Keep this independent of specific table behaviors so all resize plugins share the same lifecycle semantics.",
        ),
        (
            "src/features/editor/table/toolbar/TableControls.tsx",
            "Contextual table toolbar and typed toolbar-state selector.",
            "Shows valid commands, active headers, and border flags; routes every action through safe command wrappers.",
            "useEditorState, ToolbarPrimitives, tableCommands, tableBorders",
            "Do not call raw editor table commands here; capability and action paths must stay consistent.",
        ),
        (
            "src/features/editor/table/toolbar/TableCreationPicker.tsx",
            "Accessible table-size picker.",
            "Provides grid and manual size selection in a Floating UI dialog and clamps manual values.",
            "@floating-ui/react, React state",
            "The final insertion still goes through insertTable; UI validation is not the sole safety boundary.",
        ),
        (
            "src/features/editor/table/toolbar/TableIcons.tsx",
            "Table-specific inline SVG icon components.",
            "Provides consistent icon size and accessible decorative SVG markup.",
            "ToolbarPrimitives ICON_SIZE",
            "Keep icons presentation-only; behavior belongs in controls and commands.",
        ),
        (
            "src/features/editor/table/toolbar/index.ts",
            "Public toolbar exports.",
            "Exports TableControls and TableCreationPicker.",
            "Table toolbar components",
            "Keep the public surface small.",
        ),
        (
            "src/features/editor/components/toolbar/EditorToolbar.tsx",
            "Main editor toolbar integration point.",
            "Renders TableCreationPicker and TableControls; routes creation through insertTable.",
            "Table toolbar exports and tableCommands",
            "The contextual table toolbar is rendered below the global formatting toolbar.",
        ),
        (
            "src/features/editor/styles/tiptap-table.css",
            "All visual table, selection, border, and interaction-handle styling.",
            "Uses persistent data attributes and :is(.tiptap, .ProseMirror) selectors.",
            "Custom table attributes and plugin-generated classes",
            "CSS selectors are part of the rendering contract; update alongside attribute/class changes.",
        ),
        (
            "src/app/globals.css",
            "Global stylesheet entry point.",
            "Imports tiptap-table.css so table rendering applies throughout the editor route.",
            "Next.js App Router global CSS",
            "The local Next.js 16 guidance permits global CSS from the app tree; keep import ordering predictable.",
        ),
    ]
    for ref in file_refs:
        add_file_reference(document, *ref)

    add_heading(document, "Test Files", level=2)
    test_refs = [
        (
            "src/features/editor/table/commands/tableCommands.test.ts",
            "Unit/integration coverage for safe commands, invalid states, header preservation, selected rectangles, borders, undo/redo, and row reordering.",
        ),
        (
            "src/features/editor/table/extensions/tableExtensions.integration.test.ts",
            "Editor-level coverage for serialization, drag handle behavior, merge/split, selected-cell deletion, read-only mode, and move transactions.",
        ),
        (
            "src/features/editor/table/toolbar/TableControls.test.ts",
            "Toolbar-state coverage for inactive editors, full-table selection capability rules, and persisted border synchronization.",
        ),
        (
            "src/features/editor/table/resizing/tableDimensions.test.ts",
            "Pure dimension normalization plus row preview and outer-edge hit-detection coverage.",
        ),
        (
            "src/features/editor/table/utils/tableBorders.test.ts",
            "Border defaults and rendered data-attribute coverage.",
        ),
        (
            "src/features/editor/table/utils/mouseDragSession.test.ts",
            "Movement forwarding, one-time commit, blur cancellation, and idempotent explicit cancellation.",
        ),
        (
            "tests/integration/table-editor.spec.ts",
            "Browser-level behavior for outer/internal resizing, row resizing, borders, table block movement, horizontal positioning, focus, and portal toolbar behavior.",
        ),
    ]
    add_table(
        document,
        ["Test file", "Coverage responsibility"],
        [[a, b] for a, b in test_refs],
        [3900, 5460],
        first_col_fill=LIGHTER_BLUE,
        font_size=8.7,
    )

    add_heading(document, "13. Edge-Case Behavior Matrix", level=1)
    edge_rows = [
        ["No active editor", "Creation/action helpers return false; contextual controls are inactive.", "Safe wrapper validation"],
        ["Destroyed editor", "Commands and toolbar state return false/inactive without touching view/state.", "editor.isDestroyed checks"],
        ["Read-only editor", "Mutation controls hidden/inactive; resize and drag sessions cancel; handles become inert.", "isEditable/view.editable checks"],
        ["No selected table", "Table actions return false and TableControls is hidden.", "getActiveTable + isActive('table')"],
        ["Invalid/malformed table map", "TableMap access is caught and action returns false.", "getTableMap / plugin map helpers"],
        ["Insert row above header", "New first row is header; former header row is demoted.", "Edge header normalization"],
        ["Delete header row", "Next row is promoted to header.", "Snapshot header intent + normalization"],
        ["Delete header column", "New first column is promoted.", "Snapshot header intent + normalization"],
        ["Reorder first row", "Header semantics remain on first row/column.", "Full-table normalization"],
        ["Select every row/column", "Delete row/column toolbar actions are disabled; commands return false.", "selectedRect guard"],
        ["Delete all selected cells", "Cell content is cleared; table remains.", "deleteCellSelection shortcut override"],
        ["Merge/split invalid selection", "Capability is false and action returns false.", "canRunTableCommand / action wrapper"],
        ["Resize with invalid numeric values", "Values fall back to safe width/offset/minimum height.", "tableDimensions normalization"],
        ["Pointer beside outer edge but outside table height", "Outer resize does not activate.", "isNearTableRightEdge vertical bounds"],
        ["Table DOM disappears during drag", "Session cancels and stored state is restored.", "DOM lookup + cancel path"],
        ["Window loses focus during mouse drag", "Resize cancels and listeners are removed.", "mouseDragSession blur handling"],
        ["Table moved/deleted during plugin state", "Stored table position maps or becomes null.", "mapTablePos"],
        ["Drop inside dragged table", "Move transaction returns null.", "createTableMoveTransaction validation"],
        ["Horizontal offset exceeds space", "Offset clamps to 100 - width.", "clampTableOffsetPct"],
    ]
    add_table(
        document,
        ["Edge case", "Expected behavior", "Mechanism"],
        edge_rows,
        [2300, 4450, 2610],
        font_size=8.3,
    )

    add_heading(document, "14. Hardening Improvements", level=1)
    add_heading(document, "14.1 Command and toolbar safety", level=2)
    add_bullets(
        document,
        [
            "Introduced one typed safe-command module for all table UI mutations.",
            "Added null, destroyed-editor, read-only, no-table, invalid-map, and invalid-input guards.",
            "Wrapped underlying commands in false-returning exception boundaries.",
            "Made toolbar capability checks selection-aware for full-table row/column selections.",
            "Routed table insertion and border updates through the same safety boundary.",
        ],
    )
    add_heading(document, "14.2 Header correctness", level=2)
    add_bullets(
        document,
        [
            "Preserved header intent across insertion, deletion, undo, redo, and row reordering.",
            "Demoted displaced former header rows/columns instead of allowing duplicate header edges.",
            "Added full-table normalization for row reordering.",
            "Added focused tests for first-row/header-row and first-column/header-column behavior.",
        ],
    )
    add_heading(document, "14.3 Selection and destructive behavior", level=2)
    add_bullets(
        document,
        [
            "Overrode Tiptap's full-cell-selection table deletion shortcut.",
            "Delete and Backspace now clear selected cell content without deleting the grid.",
            "Kept explicit whole-table deletion as a deliberate toolbar action.",
        ],
    )
    add_heading(document, "14.4 Resize and drag reliability", level=2)
    add_bullets(
        document,
        [
            "Protected width/offset clamping from NaN and infinite values.",
            "Restricted outer resize hit detection to the actual vertical table bounds.",
            "Added guarded commit transactions for row height, outer width, and horizontal offset.",
            "Validated top-level table positions before moving a table node.",
            "Changed vertical drop placement to use target-block midpoint geometry.",
        ],
    )
    add_heading(document, "14.5 Readability and modularity", level=2)
    add_bullets(
        document,
        [
            "Kept responsibilities in dedicated extensions, commands, plugins, resizing helpers, DOM helpers, toolbar components, and utilities.",
            "Consolidated repeated table CSS selectors with :is(...).",
            "Added explicit toolbar-state typing and a complete inactive state.",
            "Retained the existing table folder structure because it already reflects clean ownership boundaries.",
        ],
    )

    add_heading(document, "15. Test and Verification Strategy", level=1)
    add_heading(document, "15.1 Verification completed", level=2)
    add_table(
        document,
        ["Check", "Result", "What it verifies"],
        [
            ["npm test", "40 passed", "Commands, extensions, toolbar state, utilities, resizing math, selection behavior, serialization, and plugin helpers."],
            ["npm run test:integration", "10 passed", "Real browser behavior for resize, borders, row height, table movement, offsets, toolbar focus, and portal UI."],
            ["npm run lint", "Passed", "ESLint and Next.js lint rules."],
            ["npx tsc --noEmit", "Passed", "Type correctness across the project."],
            ["npm run build", "Passed", "Next.js 16.2.7 production compilation, TypeScript stage, and static generation."],
            ["In-app browser inspection", "Passed", "New table rendered with header row, drag handle, toolbar, width/offset, and border attributes."],
        ],
        [2300, 1300, 5760],
        font_size=9.0,
    )
    add_heading(document, "15.2 Test layering", level=2)
    add_para(
        document,
        "Pure math and attribute behavior is tested at the utility layer. Command behavior is tested "
        "with a real Tiptap Editor in jsdom. Plugin integration and serialization are tested with "
        "the editor and rendered DOM. Geometry-sensitive behavior is finally verified in Playwright "
        "against the running Next.js application.",
    )
    add_callout(
        document,
        "Testing principle",
        "A command test proves document semantics. A browser test proves pointer geometry and rendered "
        "behavior. Both are required for resize and drag features.",
        "success",
    )

    add_heading(document, "16. Remaining Risks and Production Concerns", level=1)
    add_table(
        document,
        ["Concern", "Current status", "Recommended follow-up"],
        [
            ["Row reorder UI", "Safe reorderTableRow API exists, but no user-facing row reorder control is exposed.", "Add UI only when product behavior and accessibility are defined."],
            ["Merged-cell row resize", "TableMap and visited-cell logic protects persistence, but complex rowspan layouts remain browser-layout dependent.", "Add browser cases with rowspans and colspans before expanding merged-cell workflows."],
            ["Input modalities", "Custom resize and drag interactions are mouse-oriented.", "Add Pointer Events/touch support and keyboard-accessible alternatives if mobile/tablet use is required."],
            ["Browser coverage", "Automated browser tests currently target Chrome.", "Add Firefox/WebKit runs in CI if cross-browser support is a production requirement."],
            ["Very large tables", "Manual insertion allows up to 100 rows by 20 columns.", "Measure editor/update performance and consider stricter product limits or virtualization strategies."],
            ["Imported malformed HTML", "Commands fail safely, but malformed tables may remain non-editable until fixed.", "Consider an explicit repair/import validation flow using fixTables and user feedback."],
            ["Border semantics with spanning cells", "CSS edge selectors use first/last child visual positions.", "Add dedicated browser tests for complex rowspan/colspan border combinations."],
            ["Collaboration/concurrent editing", "Position mapping is defensive inside local transactions.", "Re-evaluate plugin session cancellation and mapped positions before enabling collaborative editing."],
            ["Read-only initialization", "Tiptap column resizing plugins are configured based on initial editor editability.", "Verify toggling editable mode repeatedly if read-only/edit transitions become frequent."],
        ],
        [2350, 3200, 3810],
        font_size=8.5,
    )
    add_callout(
        document,
        "Production posture",
        "The table subsystem now fails safely and has strong local/browser coverage. The primary "
        "remaining risks are advanced layout combinations, non-mouse accessibility, cross-browser "
        "geometry, and performance at product-limit table sizes.",
        "warning",
    )

    add_heading(document, "17. Safe Extension and Maintenance Guide", level=1)
    add_heading(document, "17.1 Adding a new table attribute", level=2)
    add_numbered(
        document,
        [
            "Define the attribute and backward-compatible default in the appropriate custom extension or utility.",
            "Normalize unknown/imported values before using them.",
            "Render a stable HTML/data-attribute contract if CSS or export depends on it.",
            "Apply the attribute in KnowledgeBaseTableView when the editable DOM needs explicit synchronization.",
            "Route toolbar updates through a safe command helper.",
            "Add JSON, HTML, restored-editor, and browser-rendering tests.",
        ],
    )
    add_heading(document, "17.2 Adding a new structural command", level=2)
    add_numbered(
        document,
        [
            "Add a typed command name or dedicated wrapper in tableCommands.ts.",
            "Validate editor, editability, active table, selection, and command-specific inputs.",
            "Determine whether the command can disturb edge-header semantics.",
            "Keep all related document mutations inside one transaction/history action.",
            "Expose capability and action paths through TableControls without calling raw commands.",
            "Test invalid state, first-row/header behavior, selected-cell behavior, and undo/redo.",
        ],
    )
    add_heading(document, "17.3 Adding a geometry-heavy interaction", level=2)
    add_numbered(
        document,
        [
            "Keep persistent values on nodes and live previews in DOM/plugin session state.",
            "Use tableDom helpers for position and DOM resolution.",
            "Use startMouseDragSession or a pointer-equivalent lifecycle with guaranteed cleanup.",
            "Map stored positions through transactions and cancel if the target disappears.",
            "Restore preview state before committing the final transaction.",
            "Use closeHistory so the interaction becomes one undoable action.",
            "Test cancellation, blur, read-only transition, invalid DOM, undo, and browser geometry.",
        ],
    )
    add_heading(document, "17.4 Review checklist", level=2)
    add_bullets(
        document,
        [
            "No table UI calls raw mutation commands when a safe wrapper exists.",
            "Commands return false for invalid state and do not throw.",
            "Header-row and header-column intent survives structural changes.",
            "Full-table selections do not expose misleading row/column deletion controls.",
            "Selected-cell Delete/Backspace preserves the table.",
            "Preview DOM state is always cleaned up.",
            "Persistent attributes serialize through JSON and HTML.",
            "Toolbar state updates with selection and attribute changes.",
            "Unit, editor integration, browser integration, TypeScript, lint, and build checks pass.",
        ],
    )

    add_heading(document, "Appendix A. Command and Attribute Quick Reference", level=1)
    add_table(
        document,
        ["User intent", "Preferred API", "Persistent mutation"],
        [
            ["Create table", "insertTable(editor, rows, cols)", "Inserts a table with withHeaderRow: true."],
            ["Insert/delete row or column", "runTableStructureCommand", "Structural change plus header normalization."],
            ["Merge/split cells", "runTableActionCommand", "Tiptap merge/split transaction."],
            ["Delete table", "runTableActionCommand(editor, 'deleteTable')", "Explicit table deletion."],
            ["Toggle header row/column", "runTableActionCommand", "Changes cell node types through Tiptap."],
            ["Update borders", "updateTableBorders", "Updates table border attributes."],
            ["Move row", "reorderTableRow", "moveTableRow plus full header normalization."],
            ["Resize outer table", "TableOuterResizePlugin", "Updates tableWidthPct and possibly tableOffsetPct."],
            ["Resize row", "RowResizePlugin", "Updates rowHeight on unique participating cells."],
            ["Move/position whole table", "TableDragHandlePlugin", "Moves table node vertically or updates tableOffsetPct horizontally."],
        ],
        [2500, 3100, 3760],
        font_size=8.8,
    )
    add_heading(document, "Appendix B. Key Invariants", level=1)
    add_bullets(
        document,
        [
            "The first row and first column are the only semantic header edges.",
            "Table width is between 10 and 100 percent.",
            "Horizontal offset is between 0 and 100 minus table width.",
            "Row height is at least 20 pixels.",
            "Missing border attributes mean enabled.",
            "A resize or drag commit is one history action.",
            "A stale or invalid table/editor state returns false or null instead of throwing.",
            "Selected-cell keyboard deletion clears content rather than deleting the table.",
            "Toolbar state must describe the command that will actually run.",
        ],
    )
    add_callout(
        document,
        "End state",
        "The subsystem is organized around explicit ownership, safe command boundaries, durable node "
        "attributes, deterministic cleanup, and layered verification. Future table changes should "
        "preserve those properties.",
        "success",
    )

    document.core_properties.title = "Table Architecture and Behavior Reference"
    document.core_properties.subject = "Detailed engineering reference for the kb-frontend table subsystem"
    document.core_properties.author = "KB Frontend Engineering"
    document.core_properties.keywords = "Tiptap, ProseMirror, tables, architecture, commands, toolbar, resizing"
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()

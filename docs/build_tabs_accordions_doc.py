from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Tabs_and_Accordions_Developer_Guide.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "172B4D"
MUTED = "5B677A"
LIGHT_BLUE = "E8EEF5"
LIGHTER_BLUE = "F4F7FA"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "D7DEE8"
GREEN = "EAF5EE"
GREEN_TEXT = "276749"
GOLD = "FFF4D6"
GOLD_TEXT = "7A5A00"
RED = "FBECEC"
RED_TEXT = "9B1C1C"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_TOP_BOTTOM = 80
CELL_MARGIN_START_END = 120


def set_run_font(
    run,
    *,
    name: str = "Calibri",
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    return run


def style_paragraph(
    paragraph,
    *,
    before: float = 0,
    after: float = 6,
    line: float = 1.25,
    keep_next: bool = False,
    keep_together: bool = False,
):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.keep_with_next = keep_next
    fmt.keep_together = keep_together


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (
        ("top", CELL_MARGIN_TOP_BOTTOM),
        ("bottom", CELL_MARGIN_TOP_BOTTOM),
        ("start", CELL_MARGIN_START_END),
        ("end", CELL_MARGIN_START_END),
    ):
        child = tc_mar.find(qn(f"w:{tag}"))
        if child is None:
            child = OxmlElement(f"w:{tag}")
            tc_mar.append(child)
        child.set(qn("w:w"), str(value))
        child.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]):
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError(f"Table widths must sum to {CONTENT_WIDTH_DXA}: {widths_dxa}")

    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_paragraph_box(paragraph, fill: str, border_color: str):
    p_pr = paragraph._p.get_or_add_pPr()
    shading = p_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        p_pr.append(shading)
    shading.set(qn("w:fill"), fill)

    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        border = borders.find(qn(f"w:{edge}"))
        if border is None:
            border = OxmlElement(f"w:{edge}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "8")
        border.set(qn("w:space"), "6")
        border.set(qn("w:color"), border_color)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr)
    run._r.append(fld_char_2)
    set_run_font(run, size=9, color=MUTED)


def configure_document(document: Document):
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.line_spacing = 1.1

    code = styles.add_style("Code Block", 1)
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    code._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    code.font.size = Pt(8.5)
    code.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    code.paragraph_format.space_before = Pt(3)
    code.paragraph_format.space_after = Pt(7)
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.right_indent = Inches(0.18)
    code.paragraph_format.line_spacing = 1.05

    header = section.header.paragraphs[0]
    style_paragraph(header, after=0, line=1.0)
    set_run_font(
        header.add_run("KB Frontend Tabs and Accordions"),
        size=9,
        color=MUTED,
        bold=True,
    )
    set_run_font(header.add_run("  |  Developer Guide"), size=9, color=MUTED)
    add_page_number(section.footer.paragraphs[0])


def add_heading(document: Document, text: str, level: int = 1):
    return document.add_heading(text, level=level)


def add_para(
    document: Document,
    text: str = "",
    *,
    bold_label: str | None = None,
    after: float = 6,
    line: float = 1.25,
    keep_together: bool = False,
):
    p = document.add_paragraph()
    style_paragraph(p, after=after, line=line, keep_together=keep_together)
    if bold_label and text.startswith(bold_label):
        set_run_font(p.add_run(bold_label), bold=True, color=DARK_BLUE)
        set_run_font(p.add_run(text[len(bold_label) :]), color=INK)
    else:
        set_run_font(p.add_run(text), color=INK)
    return p


def add_bullets(document: Document, items: list[str]):
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        style_paragraph(p, after=4, line=1.25)
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        set_run_font(p.add_run(item), color=INK)


def add_numbered(document: Document, items: list[str]):
    for item in items:
        p = document.add_paragraph(style="List Number")
        style_paragraph(p, after=4, line=1.25)
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        set_run_font(p.add_run(item), color=INK)


def add_callout(document: Document, title: str, body: str, kind: str = "info"):
    palette = {
        "info": (LIGHTER_BLUE, DARK_BLUE),
        "success": (GREEN, GREEN_TEXT),
        "warning": (GOLD, GOLD_TEXT),
        "risk": (RED, RED_TEXT),
    }
    fill, title_color = palette[kind]
    p = document.add_paragraph()
    set_paragraph_box(p, fill, title_color)
    style_paragraph(p, before=2, after=8, line=1.15, keep_together=True)
    set_run_font(p.add_run(f"{title}: "), bold=True, color=title_color)
    set_run_font(p.add_run(body), color=INK)


def add_code(document: Document, code: str):
    p = document.add_paragraph(style="Code Block")
    set_paragraph_box(p, LIGHTER_BLUE, BLUE)
    style_paragraph(p, before=2, after=8, line=1.05, keep_together=True)
    set_run_font(p.add_run(code), name="Consolas", size=8.5, color=DARK_BLUE)


def add_table(
    document: Document,
    headers: list[str],
    rows: list[list[str]],
    widths: list[int],
    *,
    font_size: float = 9.0,
    first_col_fill: str | None = None,
):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        style_paragraph(p, after=0, line=1.1)
        set_run_font(p.add_run(header), size=font_size, color=DARK_BLUE, bold=True)
    repeat_table_header(table.rows[0])

    for values in rows:
        row = table.add_row()
        for i, value in enumerate(values):
            cell = row.cells[i]
            if i == 0 and first_col_fill:
                set_cell_shading(cell, first_col_fill)
            p = cell.paragraphs[0]
            style_paragraph(p, after=0, line=1.12)
            set_run_font(
                p.add_run(value),
                size=font_size,
                color=DARK_BLUE if i == 0 and first_col_fill else INK,
                bold=(i == 0 and first_col_fill is not None),
            )
    set_table_geometry(table, widths)
    add_para(document, "", after=3)
    return table


def add_title_page(document: Document):
    p = document.add_paragraph()
    style_paragraph(p, before=82, after=10, line=1.0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("KB FRONTEND ENGINEERING REFERENCE"), size=10.5, color=BLUE, bold=True)

    p = document.add_paragraph()
    style_paragraph(p, after=10, line=1.0, keep_together=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(
        p.add_run("Tabs and Accordions Developer Guide"),
        size=28,
        color=INK,
        bold=True,
    )

    p = document.add_paragraph()
    style_paragraph(p, after=28, line=1.15)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(
        p.add_run(
            "Architecture, persisted schema, commands, node views, insertion flows, "
            "static rendering, edge cases, and verification"
        ),
        size=13,
        color=MUTED,
        italic=True,
    )

    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_cell_shading(table.rows[0].cells[0], LIGHT_BLUE)
    set_cell_shading(table.rows[0].cells[1], LIGHT_BLUE)
    set_run_font(
        table.rows[0].cells[0].paragraphs[0].add_run("Document detail"),
        bold=True,
        color=DARK_BLUE,
    )
    set_run_font(
        table.rows[0].cells[1].paragraphs[0].add_run("Value"),
        bold=True,
        color=DARK_BLUE,
    )
    repeat_table_header(table.rows[0])
    metadata = [
        ("Repository", r"C:\gamalearn\kb-frontend"),
        ("Primary stack", "Next.js 16.2.7, React 19.2.4, Tiptap 3.26.0, ProseMirror"),
        ("Document date", "June 11, 2026"),
        ("Scope", "All editor code and tests related to Tabs and Accordions"),
        ("Verification", "57 Vitest tests, 11 Playwright tests, lint, TypeScript, and production build"),
    ]
    for label, value in metadata:
        row = table.add_row()
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        set_run_font(row.cells[0].paragraphs[0].add_run(label), bold=True, color=DARK_BLUE)
        set_run_font(row.cells[1].paragraphs[0].add_run(value), color=INK)
    set_table_geometry(table, [2500, 6860])

    add_callout(
        document,
        "System boundary",
        "This implementation is a local Tiptap/ProseMirror feature. It does not depend on Y.js, "
        "Hocuspocus, shared cursors, or real-time collaboration assumptions.",
        "info",
    )
    document.add_page_break()


def build_document():
    document = Document()
    configure_document(document)
    add_title_page(document)

    add_heading(document, "Contents", level=1)
    add_numbered(
        document,
        [
            "Feature purpose and expected behavior",
            "Architecture overview",
            "Persisted data model and ProseMirror node structure",
            "Insertion through toolbar, slash menu, and commands",
            "Node-view behavior in editor and read-only modes",
            "Transactions, history, add/remove/reorder, and safety",
            "Static HTML rendering and export behavior",
            "File-by-file responsibility reference",
            "Edge cases and behavior matrix",
            "Testing and manual QA checklist",
            "Known limitations and future improvements",
            "Maintenance guide",
        ],
    )
    add_callout(
        document,
        "Reading guide",
        "Sections 1-7 explain behavior and architecture. Section 8 is the ownership map. "
        "Sections 9-12 are the practical review, QA, limitations, and maintenance reference.",
        "info",
    )

    add_heading(document, "1. Feature Purpose and Expected Behavior", level=1)
    add_para(
        document,
        "Tabs and Accordions let authors group rich block content without flattening it into plain "
        "text or embedding feature-specific HTML inside paragraphs. Each item body remains normal "
        "ProseMirror block content, so headings, lists, tables, links, formatting, and nested content "
        "continue to use the editor's standard schema and commands.",
    )
    add_heading(document, "Expected author behavior", level=2)
    add_bullets(
        document,
        [
            "Insert Tabs or an Accordion from the Blocks toolbar dropdown or the slash menu.",
            "Edit each tab label or accordion title inline as plain text.",
            "Add items from the container-level add button.",
            "Move or remove items from each item's action menu.",
            "Use Alt+Up and Alt+Down while editing a label/title for deliberate keyboard reordering.",
            "Edit rich content directly inside every tab card or expanded accordion panel.",
            "Undo and redo label/title edits, add, remove, reorder, and editable accordion open-state changes.",
            "Never delete the final item in a Tabs or Accordion container.",
        ],
    )
    add_heading(document, "Expected reader behavior", level=2)
    add_bullets(
        document,
        [
            "Read-only tabs render as an accessible horizontal tab list with one visible panel.",
            "Arrow Left/Right, Home, and End move and activate read-only tabs.",
            "Read-only accordions use native details/summary interaction and expose no mutation controls.",
            "Read-only interaction is local UI state and does not mutate the ProseMirror document.",
            "Static exported tabs show every panel so content remains readable without JavaScript.",
            "Static exported accordions use native details/summary markup and preserve the persisted open state.",
        ],
    )
    add_callout(
        document,
        "Important distinction",
        "Accordion open state is persisted editor data. The active read-only tab is intentionally local "
        "viewer state. Static tabs are deliberately non-interactive and show all content.",
        "success",
    )

    add_heading(document, "2. Architecture Overview", level=1)
    add_code(
        document,
        "EditorToolbar / ContentBlockPicker       SlashMenu extension\n"
        "                 \\                         /\n"
        "                  \\                       /\n"
        "                   contentBlockCommands\n"
        "                           |\n"
        "                  Tabs / Accordion schemas\n"
        "                           |\n"
        "          model factories + typed transactions\n"
        "                           |\n"
        "       container node views + item node views\n"
        "                           |\n"
        "            shared node-view DOM / item UI\n"
        "                           |\n"
        "   editor mode | read-only mode | static renderHTML\n"
        "                           |\n"
        "                    CSS + automated tests"
    )
    add_table(
        document,
        ["Layer", "Primary responsibility", "Key rule"],
        [
            ["Schema/extensions", "Define persisted nodes, attributes, parsing, static HTML, and insert commands.", "Durable content belongs in ProseMirror nodes."],
            ["Model", "Own node names, attribute types, normalization, IDs, and node/content factories.", "Unknown/imported values are normalized at boundaries."],
            ["Commands/transactions", "Validate and apply mutations as isolated history actions.", "Invalid or read-only state returns false instead of mutating."],
            ["Node views", "Render editor controls and read-only interaction around contentDOM.", "Controls stay outside editable rich content."],
            ["Toolbar/slash menu", "Discover and invoke insertion commands.", "Display metadata is shared; execution remains in commands."],
            ["Static rendering", "Serialize meaningful HTML without node-view-only controls.", "Exports remain readable without runtime JavaScript."],
            ["Styles", "Render editor, viewer, menu, and static hooks.", "Class/data-attribute contracts change with tests and docs."],
            ["Tests", "Verify document semantics, node-view behavior, accessibility, and browser workflows.", "Geometry/UI checks complement editor-level tests."],
        ],
        [1850, 4300, 3210],
        font_size=8.8,
        first_col_fill=LIGHTER_BLUE,
    )
    add_callout(
        document,
        "Separation of concerns",
        "catalog.ts contains insertion display metadata only. Command execution lives in "
        "commands/contentBlockCommands.ts. Shared node-view controls live in itemUi.ts and dom.ts. "
        "Static HTML remains in the schema extensions where Tiptap expects renderHTML.",
        "info",
    )

    add_heading(document, "3. Persisted Data Model and ProseMirror Node Structure", level=1)
    add_heading(document, "3.1 Node hierarchy", level=2)
    add_table(
        document,
        ["Node", "Content expression", "Persisted attributes", "Notes"],
        [
            ["tabs", "tabItem+", "None", "Block, defining, isolating. Requires at least one item."],
            ["tabItem", "block+", "itemId: string|null; label: string", "Defining and isolating. Body supports normal rich blocks."],
            ["accordion", "accordionItem+", "None", "Block, defining, isolating. Requires at least one item."],
            ["accordionItem", "block+", "itemId: string|null; title: string; open: boolean", "Defining and isolating. open is persisted."],
        ],
        [1650, 1750, 3350, 2610],
        font_size=8.8,
        first_col_fill=LIGHTER_BLUE,
    )
    add_heading(document, "3.2 Example JSON", level=2)
    add_code(
        document,
        '{\n'
        '  "type": "tabs",\n'
        '  "content": [\n'
        '    {\n'
        '      "type": "tabItem",\n'
        '      "attrs": { "itemId": "tab-...", "label": "Overview" },\n'
        '      "content": [{ "type": "paragraph" }]\n'
        '    }\n'
        '  ]\n'
        '}\n\n'
        '{\n'
        '  "type": "accordion",\n'
        '  "content": [\n'
        '    {\n'
        '      "type": "accordionItem",\n'
        '      "attrs": { "itemId": "accordion-...", "title": "FAQ", "open": false },\n'
        '      "content": [{ "type": "paragraph" }]\n'
        '    }\n'
        '  ]\n'
        '}'
    )
    add_heading(document, "3.3 Normalization and compatibility", level=2)
    add_bullets(
        document,
        [
            "Labels and titles are trimmed, fall back to Tab or Section when empty, and are capped at 2,000 characters.",
            "New items receive a generated itemId using crypto.randomUUID when available, with a timestamp/counter fallback.",
            "Legacy imported HTML may omit itemId, so the attribute type is nullable and read-only tabs use a positional fallback.",
            "Duplicate item IDs can exist after duplication/import. The read-only tab view disambiguates them with an occurrence suffix.",
            "Item bodies use block+, preserving rich nested content and allowing the standard editor schema to validate it.",
            "Container and item nodes are isolating, reducing accidental Backspace/Delete merges across item boundaries.",
        ],
    )

    add_heading(document, "4. Insertion Through Toolbar, Slash Menu, and Commands", level=1)
    add_heading(document, "4.1 Shared catalog", level=2)
    add_para(
        document,
        "catalog.ts is the single source for the supported block kinds and their user-facing labels "
        "and descriptions. The toolbar and slash menu read this metadata, which prevents their names "
        "and descriptions from drifting.",
    )
    add_heading(document, "4.2 Command flow", level=2)
    add_numbered(
        document,
        [
            "The UI calls insertContentBlock(editor, kind) or the slash menu calls runContentBlockInsert(chain, kind).",
            "The command boundary rejects missing, destroyed, or read-only editors.",
            "The chosen Tiptap command inserts createTabsContent() or createAccordionContent().",
            "Each default container starts with two valid items containing empty paragraphs.",
            "Tiptap places the compound block through the normal schema-aware insertContent command.",
        ],
    )
    add_heading(document, "4.3 Slash menu behavior", level=2)
    add_bullets(
        document,
        [
            "A match is recognized only in an empty TextSelection inside a non-code text block.",
            "The slash must start the text block or follow whitespace.",
            "Queries match the block kind or label prefix, such as /ta and /acc.",
            "Arrow Up/Down cycles options; Enter or Tab inserts; Escape dismisses the current match.",
            "The slash range is deleted in the same chained insertion operation.",
            "The plugin returns no decorations and handles no insertion while the editor is read-only.",
        ],
    )

    add_heading(document, "5. Node-View Behavior", level=1)
    add_heading(document, "5.1 Editable tabs", level=2)
    add_bullets(
        document,
        [
            "TabsNodeView renders a stacked editor container and an add button.",
            "Each TabItemNodeView renders a title textarea, rich contentDOM body, collapse toggle, and action menu.",
            "Every tab body remains visible by default so authors can directly edit all content.",
            "Collapsing a tab body is local editor UI state and is not persisted.",
            "Nested content edits do not rebuild title controls, preserving focus and DOM stability.",
        ],
    )
    add_heading(document, "5.2 Read-only tabs", level=2)
    add_bullets(
        document,
        [
            "TabsNodeView renders role=tablist controls and assigns role=tabpanel to runtime item DOM.",
            "aria-controls and aria-labelledby connect each control and panel.",
            "Only the active panel is visible; long labels are ellipsized and retain their full title tooltip.",
            "Arrow Left/Right wrap between tabs; Home and End jump to the first and last tab.",
            "Active state follows stable item keys when possible and remains local to the viewer.",
            "No add, move, remove, label-editing, or collapse controls are rendered.",
        ],
    )
    add_heading(document, "5.3 Editable accordions", level=2)
    add_bullets(
        document,
        [
            "AccordionNodeView renders the item list and a container-level add button.",
            "AccordionItemNodeView uses native details/summary around a rich contentDOM panel.",
            "The title textarea and item action menu are outside the rich content DOM.",
            "Native toggle events persist open only while editable and create undoable history entries.",
            "Rapid open/close updates keep the same title control and preserve the final document state.",
        ],
    )
    add_heading(document, "5.4 Read-only accordions", level=2)
    add_bullets(
        document,
        [
            "The item remains native details/summary with plain title text and a visual chevron.",
            "User toggles change only the DOM; the persisted ProseMirror open attribute is not mutated.",
            "Mutation controls and title textareas are omitted.",
        ],
    )
    add_heading(document, "5.5 Shared node-view utilities", level=2)
    add_para(
        document,
        "itemUi.ts centralizes the repeated item action controller and title/label textarea behavior. "
        "dom.ts owns lower-level DOM primitives: attribute application, shared editable-mode observers, "
        "icons, icon buttons, and keyboard-accessible action menus. This keeps Tabs and Accordions "
        "consistent without forcing their distinct container/item rendering into one abstraction.",
    )

    add_heading(document, "6. Transactions, History, and Safety", level=1)
    add_heading(document, "6.1 Transaction rules", level=2)
    add_bullets(
        document,
        [
            "Every node-view mutation checks view.editable before changing the document.",
            "resolveNodeViewPosition safely handles stale or detached node-view positions.",
            "getItemContext validates both the expected container type and item type.",
            "updateNodeAttributes is generic over the item node name and accepts only that item's typed attributes.",
            "appendItem validates the parent/item type pair before inserting.",
            "removeItem refuses to remove the final child.",
            "moveItem swaps only the adjacent pair, avoiding a full-container replacement and preserving unaffected controls.",
            "closeHistory starts each deliberate control action as a separate undoable history event.",
        ],
    )
    add_heading(document, "6.2 Reorder and deletion safety", level=2)
    add_bullets(
        document,
        [
            "Plain arrow keys inside title/label textareas move the caret and never reorder.",
            "Only Alt+Up/Alt+Down and explicit action-menu commands reorder items.",
            "Move Up/Down buttons are disabled at boundaries and refresh whenever the menu opens.",
            "The final Remove action is disabled and the transaction layer enforces the same invariant.",
            "Isolating item nodes protect boundaries from accidental Backspace/Delete merging.",
            "Reorder, add, remove, labels/titles, and editable accordion open state participate in undo/redo.",
        ],
    )
    add_callout(
        document,
        "Defense in depth",
        "Disabled buttons improve the UI, but transaction helpers independently validate editability, "
        "position, parent type, item type, destination, and minimum child count.",
        "success",
    )

    add_heading(document, "7. Static HTML Rendering and Export Behavior", level=1)
    add_heading(document, "7.1 Tabs", level=2)
    add_code(
        document,
        '<div class="kb-tabs" data-kb-tabs>\n'
        '  <section class="kb-tabs__static-item" data-kb-tab-item aria-label="Overview">\n'
        '    <h3 data-kb-tab-label-static>Overview</h3>\n'
        '    <div data-kb-tab-panel>...</div>\n'
        '  </section>\n'
        '</div>'
    )
    add_para(
        document,
        "Static tab HTML intentionally renders every section. It does not emit role=tab or "
        "role=tabpanel because no runtime script is present to implement tab interaction. This keeps "
        "the export truthful, accessible, printable, and readable in email or server-rendered output.",
    )
    add_heading(document, "7.2 Accordions", level=2)
    add_code(
        document,
        '<div class="kb-accordion" data-kb-accordion>\n'
        '  <details class="kb-accordion__item" data-kb-accordion-item open>\n'
        '    <summary class="kb-accordion__summary" data-kb-accordion-title-static>FAQ</summary>\n'
        '    <div class="kb-accordion__panel" data-kb-accordion-panel>...</div>\n'
        '  </details>\n'
        '</div>'
    )
    add_para(
        document,
        "Static accordions use native details/summary and preserve the persisted open attribute. "
        "The static summary and panel carry the same styling hooks used by the runtime presentation. "
        "External export surfaces should include tiptap-content-blocks.css when matching presentation is required.",
    )
    add_heading(document, "7.3 Parsing and round-trip behavior", level=2)
    add_bullets(
        document,
        [
            "parseHTML uses data-kb-tabs/data-kb-tab-item and data-kb-accordion/data-kb-accordion-item hooks.",
            "contentElement points to the panel wrapper, keeping labels/titles outside rich item content.",
            "Labels/titles can be recovered from data attributes or static heading/summary text.",
            "Accordion open is recovered from the native open attribute.",
            "Editor getHTML() and Tiptap generateHTML() produce the same static output in tests.",
            "Round-tripping static HTML restores rich nested content and persisted attributes.",
        ],
    )

    add_heading(document, "8. File-by-File Responsibility Reference", level=1)
    files = [
        ["src/features/editor/contentBlocks/model.ts", "Schema-adjacent model", "Node names, typed attributes, normalization, IDs, JSON/node factories."],
        ["src/features/editor/contentBlocks/catalog.ts", "Insertion catalog", "Supported block kinds plus shared labels/descriptions."],
        ["src/features/editor/contentBlocks/commands/contentBlockCommands.ts", "Command boundary", "Typed chain execution and safe toolbar/API insertion."],
        ["src/features/editor/contentBlocks/transactions.ts", "Document mutations", "Typed attribute updates, append/remove/reorder, history separation, validation."],
        ["src/features/editor/contentBlocks/extensions/Tabs.ts", "Tabs schema", "tabs/tabItem nodes, parsing, static HTML, insertTabs, node-view registration."],
        ["src/features/editor/contentBlocks/extensions/Accordion.ts", "Accordion schema", "accordion/item nodes, open/title attributes, static HTML, insertAccordion."],
        ["src/features/editor/contentBlocks/extensions/SlashMenu.ts", "Slash integration", "Slash matching, option decoration, keyboard selection, insertion."],
        ["src/features/editor/contentBlocks/extensions/index.ts", "Extension registration", "Exports the complete content-block extension list."],
        ["src/features/editor/contentBlocks/nodeViews/dom.ts", "DOM primitives", "Editable observer, icons/buttons, dynamic action menus, menu keyboard navigation."],
        ["src/features/editor/contentBlocks/nodeViews/itemUi.ts", "Shared item UI", "Label/title textarea behavior and move/remove action controller."],
        ["src/features/editor/contentBlocks/nodeViews/TabsNodeView.ts", "Tabs container view", "Editable stacked container, read-only tablist/panels, add, active state."],
        ["src/features/editor/contentBlocks/nodeViews/TabItemNodeView.ts", "Tab item view", "Inline label editing, body collapse, actions, editor/read-only item DOM."],
        ["src/features/editor/contentBlocks/nodeViews/AccordionNodeView.ts", "Accordion container view", "Item contentDOM and editable add control."],
        ["src/features/editor/contentBlocks/nodeViews/AccordionItemNodeView.ts", "Accordion item view", "details/summary, inline title editing, persisted editable open state, actions."],
        ["src/features/editor/contentBlocks/toolbar/ContentBlockPicker.tsx", "Toolbar UI", "Renders catalog options and calls safe insertion command."],
        ["src/features/editor/contentBlocks/toolbar/index.ts", "Toolbar export", "Public export for ContentBlockPicker."],
        ["src/features/editor/extensions/index.tsx", "Editor extension assembly", "Adds contentBlockExtensions to every editor/viewer Tiptap instance."],
        ["src/features/editor/components/toolbar/EditorToolbar.tsx", "Toolbar integration", "Places ContentBlockPicker in the main editor toolbar."],
        ["src/features/editor/components/KnowledgeBaseEditor.tsx", "Editable host", "Creates editable Tiptap instance and switches editability."],
        ["src/features/editor/components/KnowledgeBaseViewer.tsx", "Read-only host", "Creates read-only Tiptap instance and updates viewer content."],
        ["src/features/editor/styles/tiptap-content-blocks.css", "Rendering contract", "Editor cards, read-only tabs, accordions, static hooks, menus, slash menu."],
        ["src/app/globals.css", "Global style entry", "Imports tiptap-content-blocks.css for the app."],
        ["src/features/editor/contentBlocks/contentBlocks.integration.test.ts", "Editor-level tests", "Schema, static HTML, node views, safety, history, read-only, slash behavior."],
        ["tests/integration/content-blocks-editor.spec.ts", "Browser test", "Toolbar/slash insertion, inline editing, add controls, nested rich editing."],
        ["docs/tabs-and-accordions.md", "Short-form reference", "Concise behavior and edge-case notes."],
    ]
    add_table(
        document,
        ["File", "Ownership", "What it does"],
        files,
        [4050, 1750, 3560],
        font_size=7.9,
        first_col_fill=LIGHTER_BLUE,
    )

    add_heading(document, "9. Edge Cases and Behavior Matrix", level=1)
    edge_rows = [
        ["Missing/destroyed editor", "Safe insertion returns false.", "Command boundary"],
        ["Read-only editor", "No insert/add/move/remove/title mutation; viewer interaction remains local.", "Commands, transactions, node views"],
        ["Final item removal", "Remove is disabled and transaction refuses deletion.", "Dynamic menu + removeItem"],
        ["Move at first/last boundary", "Move command is disabled and returns false.", "Dynamic menu + moveItem"],
        ["Nested rich content edit", "Body changes without rebuilding label/title controls.", "Selective node-view update"],
        ["Long labels/titles", "Up to 2,000 chars; editor wraps; read-only tab pill truncates with tooltip.", "Normalization + CSS"],
        ["Empty/whitespace title", "Falls back to Tab or Section.", "normalizeItemLabel"],
        ["Legacy item without ID", "Loads with nullable ID and positional viewer fallback.", "Schema defaults + readContentBlockItemId"],
        ["Duplicate item IDs", "Read-only tab keys add occurrence suffixes.", "describeTabs"],
        ["Ordinary arrow keys in label", "Caret moves; no reorder.", "Textarea key handling"],
        ["Backspace/Delete at item boundary", "Items remain intact.", "Isolating node schema"],
        ["Rapid accordion toggles", "Final open state persists; title control remains stable.", "Selective update + typed attribute transaction"],
        ["Read-only accordion toggle", "DOM toggles; document open attribute stays unchanged.", "Editable guard"],
        ["Editable/read-only mode switch", "Mutation chrome is added/removed without changing content.", "Shared MutationObserver"],
        ["Static tabs without JavaScript", "All panels remain visible and semantically labeled.", "renderHTML contract"],
        ["Static accordion open state", "Native open attribute is emitted and parsed.", "renderHTML/parseHTML"],
        ["Undo/redo after controls", "Label/title, add, remove, reorder, and open actions restore correctly.", "closeHistory transactions"],
    ]
    add_table(
        document,
        ["Case", "Expected behavior", "Mechanism"],
        edge_rows,
        [2450, 4290, 2620],
        font_size=8.2,
    )

    add_heading(document, "10. Testing and Manual QA Checklist", level=1)
    add_heading(document, "10.1 Automated verification completed", level=2)
    add_table(
        document,
        ["Check", "Result", "Coverage"],
        [
            ["npm test", "57 passed", "All Vitest suites; 17 focused Tabs/Accordions scenarios."],
            ["npm run test:integration", "11 passed", "Full Playwright browser suite, including content-block workflow."],
            ["npm run lint", "Passed", "ESLint and Next.js rules; rerun sequentially after Playwright."],
            ["npx tsc --noEmit", "Passed", "Strict TypeScript project validation."],
            ["npm run build", "Passed", "Next.js 16.2.7 optimized production build and static page generation."],
        ],
        [2400, 1500, 5460],
        font_size=8.8,
        first_col_fill=LIGHTER_BLUE,
    )
    add_heading(document, "10.2 Focused automated scenarios", level=2)
    add_bullets(
        document,
        [
            "Valid insertion and read-only refusal.",
            "Rich nested content and static HTML round trip.",
            "Inline label/title editing and live DOM metadata synchronization.",
            "Add, remove, adjacent-only reorder, minimum-one-item enforcement, and dynamic disabled states.",
            "Ordinary arrow keys and Backspace/Delete boundary safety.",
            "Stable controls during nested edits and unaffected controls during reorder.",
            "Undo/redo for labels, add, remove, reorder, and accordion open state.",
            "Long labels/titles, rapid accordion toggles, read-only tabs/accordions, mode switching, legacy/duplicate IDs.",
            "Slash-menu keyboard insertion and item action-menu keyboard navigation.",
        ],
    )
    add_heading(document, "10.3 Manual QA checklist", level=2)
    add_bullets(
        document,
        [
            "Insert Tabs from the toolbar and Accordion from /accordion; repeat with the opposite insertion paths.",
            "Rename labels/titles with short, empty, whitespace-only, and very long text.",
            "Confirm Enter commits, Escape restores, plain arrows move the caret, and Alt+Up/Down reorders.",
            "Add several items, move first/middle/last items, and verify content remains paired with its label/title.",
            "Attempt to move beyond boundaries and remove the final item.",
            "Place headings, lists, links, formatting, and tables inside item bodies; edit and undo/redo them.",
            "Collapse editable tab bodies and confirm collapse is local UI state.",
            "Rapidly toggle editable accordions, then undo/redo open-state changes.",
            "Use action menus by mouse and keyboard, including Arrow Up/Down, Home, End, and Escape.",
            "Switch an existing editor between editable and read-only; confirm mutation controls disappear/reappear.",
            "In read-only tabs, verify click, Arrow Left/Right, Home, End, focus order, visible panel, and long-label tooltip.",
            "In read-only accordions, toggle items and confirm persisted JSON is unchanged.",
            "Export HTML and verify every tab panel is visible, accordion open is preserved, and no editor controls appear.",
            "Reload exported HTML and confirm labels/titles, rich content, IDs, and open state round-trip.",
            "Inspect narrow/mobile widths for wrapping, overflow, menus, and long-label behavior.",
        ],
    )

    add_heading(document, "11. Known Limitations and Future Improvements", level=1)
    add_table(
        document,
        ["Area", "Current limitation", "Potential improvement"],
        [
            ["Labels/titles", "Plain-text attributes only; capped at 2,000 characters.", "Keep plain text unless product requirements justify rich label nodes."],
            ["Active tab", "Read-only active tab is local and not persisted.", "Persist only if product semantics require a default active tab."],
            ["Static tabs", "All panels show; no JavaScript tab switcher is exported.", "Add an optional external enhancer without changing semantic fallback."],
            ["Item count", "No product-level maximum.", "Add a configurable limit if UX/performance requirements demand it."],
            ["Item IDs", "Legacy items can be null; duplicated content can repeat IDs.", "Add an explicit migration/backfill command if globally unique IDs become required."],
            ["Reordering UI", "Action menu and Alt+Arrow only; no drag-and-drop.", "Add accessible drag/reorder only with keyboard parity and strong tests."],
            ["Accordion history", "Every editable open/close is a deliberate history event.", "Group or exclude toggles only if product expectations change."],
            ["External export styling", "Matching presentation requires tiptap-content-blocks.css.", "Provide a packaged export stylesheet or server-render helper."],
            ["Browser coverage", "Playwright runs the configured Chrome channel.", "Add Firefox/WebKit CI runs when required."],
            ["Viewer browser fixture", "Read-only behavior is deeply tested in jsdom, less extensively in Playwright.", "Add a dedicated viewer route/fixture for browser-level accessibility checks."],
            ["Collaboration", "No Y.js/Hocuspocus or concurrent-edit assumptions.", "Design and verify mapping/session behavior separately before collaboration work."],
        ],
        [1900, 3600, 3860],
        font_size=8.2,
    )
    add_callout(
        document,
        "Production posture",
        "The current feature is modular, typed, defensively validated, undoable, exportable, and covered "
        "at editor and browser levels. Future work should preserve the separation between persisted "
        "document state, local viewer/editor UI state, and static fallback markup.",
        "success",
    )

    add_heading(document, "12. Maintenance Guide", level=1)
    add_heading(document, "12.1 Adding a new content-block kind", level=2)
    add_numbered(
        document,
        [
            "Add its kind and user-facing metadata to catalog.ts.",
            "Add schema/model factories and typed attributes in the appropriate model/extension files.",
            "Add a command execution branch in contentBlockCommands.ts.",
            "Register the extension in contentBlocks/extensions/index.ts.",
            "Keep node-view-only UI under nodeViews and static HTML in renderHTML.",
            "Add toolbar/slash, read-only, static export, round-trip, history, and safety tests.",
        ],
    )
    add_heading(document, "12.2 Changing Tabs or Accordion behavior", level=2)
    add_bullets(
        document,
        [
            "Do not put display metadata and command execution in the same module.",
            "Do not mutate document state directly from UI without a typed transaction helper.",
            "Do not rely only on disabled UI; validate again at the transaction boundary.",
            "Avoid rebuilding item controls for ordinary nested rich-content changes.",
            "Preserve contentDOM ownership and keep mutation chrome outside editable content.",
            "Keep editor, read-only, and static-export semantics explicit and separately tested.",
            "Update CSS selectors, static classes/data attributes, tests, and this guide together.",
            "Do not introduce Y.js, Hocuspocus, or collaboration behavior implicitly.",
        ],
    )
    add_heading(document, "12.3 Review checklist for future changes", level=2)
    add_bullets(
        document,
        [
            "Schema still enforces at least one item and rich block content.",
            "Legacy/malformed imported values normalize safely.",
            "Add/remove/reorder actions are one deliberate history step and preserve content.",
            "Boundary keys cannot accidentally merge or delete items.",
            "Read-only mode renders no mutation controls and does not mutate document state.",
            "Static output is meaningful without JavaScript and round-trips through parseHTML.",
            "Action menus and tab controls remain keyboard accessible.",
            "Focused tests, full Vitest, full Playwright, lint, TypeScript, and production build pass.",
        ],
    )

    document.core_properties.title = "Tabs and Accordions Developer Guide"
    document.core_properties.subject = "Engineering reference for the kb-frontend Tabs and Accordions editor features"
    document.core_properties.author = "KB Frontend Engineering"
    document.core_properties.keywords = "Tiptap, ProseMirror, tabs, accordion, editor, node views, static HTML"
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
